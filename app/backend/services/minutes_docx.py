"""تولید فایل Word (docx) واقعی صورتجلسه با چیدمان راست‌به‌چپ فارسی.

نکته‌های طراحی:

* خروجی یک بستهٔ OOXML واقعی است (با ``python-docx``)، نه HTML با پسوند ``.docx``؛
  بنابراین در Word، WPS و LibreOffice بدون هشدار باز می‌شود.
* راست‌به‌چپ بودن در سه سطح اعمال می‌شود: بخش سند (``w:bidi``)، هر پاراگراف
  (``w:bidi`` + تراز راست) و هر ``run`` (``w:rtl``)؛ فونت هم برای متن لاتین و هم
  برای «complex script» تنظیم می‌شود تا فارسی درست شکل بگیرد.
* تاریخ‌ها همیشه شمسی و با رقم‌های فارسی نوشته می‌شوند و زمان به منطقهٔ زمانی
  سازمان تبدیل می‌گردد.
* لوگوی سازمان اگر در دسترس باشد در سربرگ می‌آید؛ نبودن لوگو خروجی را نمی‌شکند.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime, timedelta, timezone
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

# فونت پیش‌فرض؛ روی ویندوز/آفیس فارسی همیشه موجود است.
BODY_FONT = "Tahoma"

PERSIAN_DIGITS = "۰۱۲۳۴۵۶۷۸۹"

JALALI_MONTHS = (
    "فروردین",
    "اردیبهشت",
    "خرداد",
    "تیر",
    "مرداد",
    "شهریور",
    "مهر",
    "آبان",
    "آذر",
    "دی",
    "بهمن",
    "اسفند",
)

RSVP_LABELS = {
    "pending": "بی‌پاسخ",
    "accepted": "حضور می‌یابم",
    "declined": "حضور ندارم",
    "tentative": "نامطمئن",
}

ACTION_STATUS_LABELS = {
    "open": "باز",
    "in_progress": "در حال انجام",
    "done": "انجام‌شده",
    "overdue": "دارای تأخیر",
}

MINUTES_STATUS_LABELS = {
    "draft": "پیش‌نویس",
    "in_review": "در انتظار تأیید",
    "approved": "تأییدشده",
    "locked": "نهایی و قفل‌شده",
}


# ---------------------------------------------------------------------------
# تاریخ شمسی
# ---------------------------------------------------------------------------


def to_persian_digits(value: Any) -> str:
    return "".join(PERSIAN_DIGITS[int(ch)] if ch.isdigit() else ch for ch in str(value))


def gregorian_to_jalali(gy: int, gm: int, gd: int) -> tuple[int, int, int]:
    """تبدیل تاریخ میلادی به هجری شمسی (الگوریتم استاندارد بی‌خطا در بازهٔ کاربردی)."""
    month_days = (0, 31, 59, 90, 120, 151, 181, 212, 243, 273, 304, 334)
    gy2 = gy + 1 if gm > 2 else gy
    days = (
        355666
        + (365 * gy)
        + ((gy2 + 3) // 4)
        - ((gy2 + 99) // 100)
        + ((gy2 + 399) // 400)
        + gd
        + month_days[gm - 1]
    )
    jy = -1595 + (33 * (days // 12053))
    days %= 12053
    jy += 4 * (days // 1461)
    days %= 1461
    if days > 365:
        jy += (days - 1) // 365
        days = (days - 1) % 365
    if days < 186:
        jm = 1 + (days // 31)
        jd = 1 + (days % 31)
    else:
        jm = 7 + ((days - 186) // 30)
        jd = 1 + ((days - 186) % 30)
    return jy, jm, jd


def _parse_iso(value: Optional[str]) -> Optional[datetime]:
    text = (value or "").strip()
    if not text:
        return None
    if text.endswith("Z"):
        text = text[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(text)
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed


def _localize(value: datetime, tz_name: str) -> datetime:
    """تبدیل زمان UTC به منطقهٔ زمانی سازمان؛ در نبود پایگاه zoneinfo، تهران +۳:۳۰."""
    try:
        from zoneinfo import ZoneInfo

        return value.astimezone(ZoneInfo(tz_name or "Asia/Tehran"))
    except Exception:  # noqa: BLE001 - نبود tzdata نباید خروجی را بشکند
        return value.astimezone(timezone(timedelta(hours=3, minutes=30)))


def jalali_date(value: Optional[str], tz_name: str = "Asia/Tehran") -> str:
    parsed = _parse_iso(value)
    if parsed is None:
        return "—"
    local = _localize(parsed, tz_name)
    jy, jm, jd = gregorian_to_jalali(local.year, local.month, local.day)
    return f"{to_persian_digits(jd)} {JALALI_MONTHS[jm - 1]} {to_persian_digits(jy)}"


def jalali_datetime(value: Optional[str], tz_name: str = "Asia/Tehran") -> str:
    parsed = _parse_iso(value)
    if parsed is None:
        return "—"
    local = _localize(parsed, tz_name)
    jy, jm, jd = gregorian_to_jalali(local.year, local.month, local.day)
    clock = f"{local.hour:02d}:{local.minute:02d}"
    return (
        f"{to_persian_digits(jd)} {JALALI_MONTHS[jm - 1]} {to_persian_digits(jy)}"
        f" — ساعت {to_persian_digits(clock)}"
    )


# ---------------------------------------------------------------------------
# ابزارهای راست‌به‌چپ
# ---------------------------------------------------------------------------


def _set_section_rtl(section) -> None:
    pr = section._sectPr
    bidi = pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pr.append(bidi)


def _style_run(run, *, size: int, bold: bool = False, color: Optional[str] = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:cs"), BODY_FONT)
    for tag in ("w:rtl", "w:lang"):
        element = OxmlElement(tag)
        if tag == "w:lang":
            element.set(qn("w:bidi"), "fa-IR")
        rpr.append(element)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(size * 2))
    rpr.append(size_cs)


def _rtl_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))
    paragraph.alignment = alignment


def _add_paragraph(
    document,
    text: str,
    *,
    size: int = 11,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    space_after: int = 4,
    color: Optional[str] = None,
):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text or "")
    _style_run(run, size=size, bold=bold, color=color)
    _rtl_paragraph(paragraph, alignment)
    return paragraph


def _add_heading(document, text: str) -> None:
    paragraph = _add_paragraph(document, text, size=13, bold=True, space_after=2, color="1F4E9C")
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "BDDBFE")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def _set_table_rtl(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr.find(qn("w:bidiVisual")) is None:
        tbl_pr.append(OxmlElement("w:bidiVisual"))


def _shade_cell(cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _fill_cell(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text or "—")
    _style_run(run, size=size, bold=bold)
    _rtl_paragraph(paragraph)


def _add_table(document, headers: List[str], rows: List[List[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_rtl(table)
    for index, title in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, "EAF2FE")
        _fill_cell(cell, title, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            _fill_cell(cells[index], value)
    return table


# ---------------------------------------------------------------------------
# ساخت سند
# ---------------------------------------------------------------------------


def _clean_inline_markdown(text: str) -> str:
    """حذف نشانه‌های مارک‌داون درون‌خطی تا خروجی Word متن تمیز (بدون ** و ` و …) باشد."""
    cleaned = str(text or "")
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)  # بولد
    cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
    cleaned = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", cleaned)  # ایتالیک
    cleaned = cleaned.replace("**", "").replace("__", "")
    cleaned = re.sub(r"~~(.+?)~~", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)  # کد درون‌خطی
    return cleaned.strip()


def _markdown_to_lines(body: str) -> List[tuple[str, str]]:
    """تبدیل مارک‌داون صورتجلسه به دنبالهٔ (نوع، متن تمیز) برای نوشتن در Word.

    خروجی هیچ نشانهٔ مارک‌داونی ندارد: عناوین (###/##/#)، فهرست‌های گلوله‌ای و
    شماره‌دار و نشانه‌های درون‌خطی (**، *، ` و …) همگی به قالب واقعی تبدیل
    می‌شوند.
    """
    output: List[tuple[str, str]] = []
    for raw_line in (body or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("###"):
            output.append(("h3", _clean_inline_markdown(line.lstrip("#").strip())))
        elif line.startswith("##"):
            output.append(("h2", _clean_inline_markdown(line.lstrip("#").strip())))
        elif line.startswith("#"):
            output.append(("h2", _clean_inline_markdown(line.lstrip("#").strip())))
        elif line.startswith(("- ", "* ", "• ")):
            output.append(("li", _clean_inline_markdown(line[2:].strip())))
        else:
            ordered = re.match(r"^(\d+)[.)]\s+(.*)$", line)
            if ordered:
                output.append(("oli", f"{ordered.group(1)}. {_clean_inline_markdown(ordered.group(2))}"))
                continue
            output.append(("p", _clean_inline_markdown(line)))
    return output


def safe_file_name(title: str, starts_at: Optional[str], tz_name: str) -> str:
    """نام فایل فارسی و امن (بدون نویسه‌های ممنوعهٔ سیستم فایل)."""
    cleaned = re.sub(r'[\\/:*?"<>|\r\n\t]+', " ", title or "").strip()
    cleaned = re.sub(r"\s+", " ", cleaned)[:80] or "جلسه"
    stamp = jalali_date(starts_at, tz_name).replace(" ", "-")
    if stamp == "—":
        stamp = "بدون-تاریخ"
    return f"صورتجلسه-{cleaned}-{stamp}.docx"


def build_minutes_docx(payload: Dict[str, Any], logo_bytes: Optional[bytes] = None) -> bytes:
    """ساخت بایت‌های فایل docx صورتجلسه از بستهٔ خروجی جلسه."""
    organization = payload.get("organization") or {}
    meeting = payload.get("meeting") or {}
    minutes = payload.get("minutes") or {}
    tz_name = organization.get("timezone") or "Asia/Tehran"

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    _set_section_rtl(section)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)

    # سربرگ سازمان + لوگو
    if logo_bytes:
        try:
            header_paragraph = document.add_paragraph()
            _rtl_paragraph(header_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
            header_paragraph.add_run().add_picture(io.BytesIO(logo_bytes), width=Cm(4.5))
        except Exception:  # noqa: BLE001 - لوگوی نامعتبر خروجی را نمی‌شکند
            logger.warning("درج لوگو در فایل Word ناموفق بود", exc_info=True)

    _add_paragraph(
        document,
        organization.get("name") or "",
        size=12,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    _add_paragraph(
        document,
        f"صورتجلسه: {meeting.get('title') or ''}",
        size=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    # جدول مشخصات جلسه
    status = str(minutes.get("status") or "")
    info_rows = [
        ["تاریخ و ساعت", jalali_datetime(meeting.get("starts_at"), tz_name)],
        ["مدت جلسه", f"{to_persian_digits(meeting.get('duration_minutes') or 0)} دقیقه"],
        ["نوع جلسه", meeting.get("meeting_type") or "—"],
        ["دبیر جلسه", meeting.get("secretary_name") or "—"],
        ["محل برگزاری", meeting.get("location") or meeting.get("online_url") or "—"],
        [
            "وضعیت صورتجلسه",
            (MINUTES_STATUS_LABELS.get(status, status) or "ثبت نشده")
            + (
                f" — نسخهٔ {to_persian_digits(minutes.get('current_version') or 0)}"
                if minutes
                else ""
            ),
        ],
    ]
    _add_table(document, ["عنوان", "مقدار"], info_rows)
    _add_paragraph(document, "", size=8, space_after=0)

    if minutes.get("summary"):
        _add_heading(document, "خلاصهٔ جلسه")
        _add_paragraph(document, _clean_inline_markdown(str(minutes.get("summary"))))

    # دستور جلسه
    agenda = payload.get("agenda") or []
    _add_heading(document, "دستور جلسه")
    if not agenda:
        _add_paragraph(document, "بندی ثبت نشده است.")
    else:
        _add_table(
            document,
            ["ردیف", "عنوان بند", "زمان (دقیقه)", "مسئول"],
            [
                [
                    to_persian_digits(item.get("position") or index + 1),
                    item.get("title") or "—",
                    to_persian_digits(item.get("planned_minutes") or 0),
                    item.get("owner_name") or "—",
                ]
                for index, item in enumerate(agenda)
            ],
        )

    # حاضران و غایبان
    participants = payload.get("participants") or []
    attendees = [p for p in participants if p.get("attended")]
    absentees = [p for p in participants if not p.get("attended")]
    _add_heading(document, "حاضران و غایبان")
    _add_paragraph(
        document,
        "حاضران: " + ("، ".join(p.get("full_name") or "" for p in attendees) or "ثبت نشده"),
    )
    _add_paragraph(
        document,
        "غایبان: " + ("، ".join(p.get("full_name") or "" for p in absentees) or "ندارد"),
    )
    if participants:
        _add_table(
            document,
            ["نام", "پاسخ دعوت", "حضور"],
            [
                [
                    person.get("full_name") or "—",
                    RSVP_LABELS.get(str(person.get("rsvp_status") or ""), "—"),
                    "حاضر" if person.get("attended") else "غایب",
                ]
                for person in participants
            ],
        )

    # متن صورتجلسه
    _add_heading(document, "متن صورتجلسه")
    lines = _markdown_to_lines(str(minutes.get("body_markdown") or ""))
    if not lines:
        _add_paragraph(document, "متن صورتجلسه ثبت نشده است.")
    for kind, text in lines:
        if kind == "h2":
            _add_paragraph(document, text, size=12, bold=True, space_after=2)
        elif kind == "h3":
            _add_paragraph(document, text, size=11, bold=True, space_after=2)
        elif kind in ("li", "oli"):
            _add_paragraph(document, f"• {text}" if kind == "li" else text)
        else:
            _add_paragraph(document, text)

    # مصوبات
    decisions = payload.get("decisions") or []
    _add_heading(document, "مصوبات")
    if not decisions:
        _add_paragraph(document, "مصوبه‌ای ثبت نشده است.")
    else:
        _add_table(
            document,
            ["ردیف", "مصوبه", "شرح"],
            [
                [
                    to_persian_digits(item.get("position") or index + 1),
                    _clean_inline_markdown(item.get("title") or "—"),
                    _clean_inline_markdown(item.get("description") or "—"),
                ]
                for index, item in enumerate(decisions)
            ],
        )

    # اقدامات
    actions = payload.get("actions") or []
    _add_heading(document, "اقدامات و پیگیری‌ها")
    if not actions:
        _add_paragraph(document, "اقدامی ثبت نشده است.")
    else:
        _add_table(
            document,
            ["اقدام", "مسئول", "مهلت", "وضعیت"],
            [
                [
                    _clean_inline_markdown(item.get("title") or "—"),
                    item.get("owner_name") or "—",
                    jalali_date(item.get("due_date"), tz_name),
                    ACTION_STATUS_LABELS.get(str(item.get("status") or ""), "—"),
                ]
                for item in actions
            ],
        )

    # امضاها
    _add_paragraph(document, "", size=10, space_after=10)
    _add_heading(document, "امضاها")
    signature_table = _add_table(
        document,
        ["دبیر جلسه", "رئیس جلسه", "تاریخ امضا"],
        [[" ", " ", " "], [" ", " ", " "]],
    )
    for row in signature_table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(14)
            cell.paragraphs[0].paragraph_format.space_after = Pt(14)

    if minutes.get("locked_at"):
        _add_paragraph(
            document,
            f"این صورتجلسه در {jalali_datetime(minutes.get('locked_at'), tz_name)} قفل نهایی شده است.",
            size=9,
        )
    _add_paragraph(
        document,
        "تهیه‌شده با «ویدارا - نسخه جلسات»",
        size=9,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()